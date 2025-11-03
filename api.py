# -*- coding: utf-8 -*-

# --- START: 核心代码修改处 (一次性解决所有路径问题) ---
import sys
import os

# 获取当前脚本所在目录的绝对路径
try:
    current_dir = os.path.dirname(os.path.abspath(__file__))
except NameError:
    current_dir = os.getcwd()

# 构建 chineseocr 目录的绝对路径
chineseocr_path = os.path.join(current_dir, 'chineseocr')

# 将 chineseocr 目录添加到 Python 解释器的模块搜索路径的最前面
if chineseocr_path not in sys.path:
    sys.path.insert(0, chineseocr_path)
# --- END: 核心代码修改处 ---


"""
运行前请先：
1. 安装所有依赖:
   pip install "fastapi[all]" uvicorn torch transformers python-docx sentencepiece "pycorrector==1.1.3" PyMuPDF Pillow onnxruntime numpy opencv-python pyclipper shapely
2. 运行 download_models.py 脚本，将所有模型文件下载到本地。
   - 错别字模型将位于 ./macbert-csc-local/
   - 语法纠错模型将位于 ./bart-cgec-local/
3. 确保 'chineseocr' 文件夹与本脚本位于同一目录下。
4. (可选，为实现GPU加速) 确保您已安装支持CUDA的NVIDIA显卡驱动，并安装与驱动匹配的PyTorch GPU版本。
"""

import torch
from transformers import BertTokenizer, BertForMaskedLM, BartForConditionalGeneration
import docx
import re
import difflib
import html
import pycorrector
from pycorrector import Corrector
import fitz  # PyMuPDF
from PIL import Image
import io
import base64
import tempfile
from typing import Dict, Any

from fastapi import FastAPI, File, UploadFile, HTTPException
from pydantic import BaseModel

try:
    from chineseocr.model import OcrHandle
    CHINESEOCR_AVAILABLE = True
except ImportError as e:
    print(f"警告: 无法导入 chineseocr.model.OcrHandle，PDF处理功能将不可用。错误: {e}")
    CHINESEOCR_AVAILABLE = False
    OcrHandle = None

# ----------------- 全局变量与模型加载 -----------------
print("--- 正在执行模型加载 ---")

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
print(f"将使用设备: {device}")

csc_tokenizer, csc_model, gec_tokenizer, gec_model, corrector_instance, ocr_engine = (None,) * 6
try:
    local_csc_path = "./macbert-csc-local"
    csc_tokenizer = BertTokenizer.from_pretrained(local_csc_path)
    csc_model = BertForMaskedLM.from_pretrained(local_csc_path)
    csc_model.to(device)
    csc_model.eval()
    print("【错别字】校对模型加载完成。")

    local_gec_path = "./bart-cgec-local"
    gec_tokenizer = BertTokenizer.from_pretrained(local_gec_path)
    gec_model = BartForConditionalGeneration.from_pretrained(local_gec_path)
    gec_model.to(device)
    gec_model.eval()
    print("【语法】纠错模型加载完成。")

    corrector_instance = Corrector()
    print("PyCorrector 实例初始化完成。")

    if CHINESEOCR_AVAILABLE:
        ocr_engine = OcrHandle()
        print("本地 chineseocr 引擎初始化完成。")

except Exception as e:
    print(f"FATAL: 模型加载失败，程序无法运行。错误: {e}")
    sys.exit(1)
print("--- 所有模型加载完毕 ---")


GEC_DELETABLE_WORDS = [
    '的', '了', '着', '我', '你', '他', '她', '它', '我们', '你们', '他们',
    '使', '被', '将', '把', '给', '就', '都', '而',
    '大约', '大概', '将近', '左右', '非常', '十分', '特别', '尤其'
]


# ----------------- 核心业务逻辑服务 -----------------
class CorrectionService:
    """
    封装所有文本和文档校对的核心逻辑，与Web框架解耦。
    """
    def correct_text_pipeline(self, text: str):
        """处理纯文本校对的完整流程"""
        highlight_output, final_text, logs_output = self._internal_pipeline(text)
        return highlight_output, final_text, logs_output

    def correct_document_pipeline(self, filepath: str, filename: str):
        """处理文档校对的完整流程"""
        original_text, debug_ocr_results = self._read_document(filepath, filename)

        if original_text is None:
             raise ValueError("读取文档失败或内容为空。")
        if not original_text.strip():
            raise ValueError("文档内容为空。")

        text_with_placeholder = re.sub(r'(\n\s*){2,}', '<<PARAGRAPH_BREAK>>', original_text)
        text_with_lines_merged = text_with_placeholder.replace('\n', '')
        preprocessed_text = text_with_lines_merged.replace('<<PARAGRAPH_BREAK>>', '\n\n').strip()

        highlight_result, final_text, logs_output = self._internal_pipeline(preprocessed_text)

        try:
            doc = docx.Document()
            for para in final_text.split('\n'):
                doc.add_paragraph(para)
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_bytes = doc_buffer.getvalue()
        except Exception as e:
            raise RuntimeError(f"创建Word文档失败: {e}")

        return final_text, highlight_result, doc_bytes, "\n\n".join(debug_ocr_results), logs_output

    def _read_document(self, filepath, filename):
        debug_ocr_results = []
        try:
            if filename.endswith(".txt"):
                with open(filepath, 'r', encoding='utf-8') as f:
                    return f.read(), []
            elif filename.endswith(".docx"):
                doc = docx.Document(filepath)
                return "\n".join([para.text for para in doc.paragraphs]), []
            elif filename.endswith(".pdf"):
                if not ocr_engine:
                    raise ConnectionError("本地 chineseocr 引擎未成功初始化，无法处理PDF文件。")

                all_text_parts = []
                with fitz.open(filepath) as pdf_doc:
                    total_pages = len(pdf_doc)
                    for i, page in enumerate(pdf_doc):
                        print(f"正在识别PDF第 {i + 1}/{total_pages} 页...")
                        pix = page.get_pixmap(dpi=300)
                        pil_image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        result_list = ocr_engine.text_predict(pil_image, short_size=960)

                        page_text_raw = "\n".join([item[1] for item in result_list]) if result_list else ""

                        lines = page_text_raw.split('\n')
                        cleaned_lines = []
                        for line in lines:
                            cleaned_line = re.sub(r'^\s*\d+[\.\s、\t）)]*', '', line)
                            cleaned_lines.append(cleaned_line)
                        page_text = "\n".join(cleaned_lines)

                        all_text_parts.append(page_text)
                        debug_ocr_results.append(f"--- 第 {i + 1}/{total_pages} 页识别结果 ---\n\n{page_text_raw}")
                return "\n\n".join(all_text_parts), debug_ocr_results
            else:
                raise ValueError("不支持的文件格式。")
        except Exception as e:
            raise IOError(f"读取文件时出错: {e}")

    def _internal_pipeline(self, text):
        if not text or not text.strip():
            return {"text": "", "entities": []}, "", "本次运行无审查冲突。"

        print("开始智能句子切分...")
        sentences = self._split_text_into_sentences(text)
        if not sentences:
            return {"text": text, "entities": []}, text, "本次运行无审查冲突。"

        current_run_logs, final_sentences, all_details = [], [], []
        current_offset, total_sentences = 0, len(sentences)

        for i, current_sent in enumerate(sentences):
            print(f"正在处理第 {i + 1}/{total_sentences} 句...")
            csc_suggestion = self._correct_chunk_csc(current_sent)
            sent_after_csc_filtered = self._apply_safe_correction_rules(current_sent, csc_suggestion, "CSC", current_run_logs)
            gec_suggestion = self._correct_chunk_gec(sent_after_csc_filtered)
            final_sent = self._apply_safe_correction_rules(sent_after_csc_filtered, gec_suggestion, "GEC", current_run_logs)

            if not re.search(r'[。！？；]$', final_sent.strip()) and final_sent.strip():
                final_sent += '。'

            csc_details = self._get_highlight_details(current_sent, sent_after_csc_filtered, current_offset, "错字")
            all_details.extend(csc_details)
            gec_details = self._get_highlight_details(sent_after_csc_filtered, final_sent, current_offset, "语法")
            all_details.extend(gec_details)

            final_sentences.append(final_sent)
            current_offset += len(final_sent)

        print("整合结果...")
        final_text_str = "".join(final_sentences)
        unique_details = []
        seen_ranges = set()
        for detail in sorted(all_details, key=lambda x: x['start']):
            range_tuple = (detail['start'], detail['end'], detail['entity'])
            if range_tuple not in seen_ranges:
                unique_details.append(detail)
                seen_ranges.add(range_tuple)

        highlight_output = {"text": final_text_str, "entities": unique_details}
        logs_output_str = "\n\n".join([f"--- 冲突 {i+1} ---\n{log}" for i, log in enumerate(current_run_logs)]) or "本次运行无审查冲突。"
        return highlight_output, final_text_str, logs_output_str

    def _split_text_into_sentences(self, text):
        if not text or not text.strip(): return []
        sentences = re.split(r'([。！？；\n])', text)
        return [sentences[i] + (sentences[i+1] if i+1 < len(sentences) else '') for i in range(0, len(sentences), 2) if sentences[i]]

    def _correct_chunk_csc(self, text_chunk):
        with torch.no_grad():
            inputs = csc_tokenizer([text_chunk], padding=True, return_tensors='pt', max_length=512, truncation=True)
            inputs = {k: v.to(device) for k, v in inputs.items()}
            outputs = csc_model(**inputs)
            logits, input_ids = outputs.logits[0], inputs['input_ids'][0]
            predicted_ids = torch.argmax(logits, dim=-1)
            final_ids = input_ids.clone()
            for i in range(1, len(input_ids) - 1):
                original_id, predicted_id = input_ids[i], predicted_ids[i]
                if original_id != predicted_id and (logits[i, predicted_id] - logits[i, original_id]) > 0.9:
                    final_ids[i] = predicted_id
            return re.sub(r'\s+', '', csc_tokenizer.decode(final_ids, skip_special_tokens=True))

    def _correct_chunk_gec(self, text_chunk):
        if not text_chunk.strip(): return text_chunk
        input_ids = gec_tokenizer.encode(text_chunk.lstrip(), return_tensors='pt', max_length=512, truncation=True)
        input_ids = input_ids.to(device)
        with torch.no_grad():
            outputs = gec_model.generate(input_ids, max_length=256, num_beams=5, early_stopping=True)
        bart_corrected = gec_tokenizer.decode(outputs[0], skip_special_tokens=True)
        bart_corrected_cleaned = re.sub(r'\s+|([，。！？；])\s*', r'\1', bart_corrected)
        if corrector_instance:
            py_corrected = corrector_instance.correct(bart_corrected_cleaned)['target']
            return re.sub(r'\s+|([，。！？；])\s*', r'\1', py_corrected)
        return bart_corrected_cleaned

    def _apply_safe_correction_rules(self, original_sent, corrected_sent, stage, logs_list):
        stage_map = {"CSC": "错别字审查", "GEC": "语法审查"}
        tag_map = {"replace": "替换", "delete": "删除", "insert": "插入", "equal": "相同"}

        if original_sent == corrected_sent: return original_sent
        corrected_sent = re.sub(r'\s+', '', corrected_sent)
        final_parts = []
        matcher = difflib.SequenceMatcher(None, original_sent, corrected_sent)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            original_fragment, corrected_fragment = original_sent[i1:i2], corrected_sent[j1:j2]
            if tag == 'equal':
                final_parts.append(original_fragment)
                continue
            is_whitelisted = (stage == "GEC" or (len(original_fragment) == 1 and len(corrected_fragment) == 1) or set([original_fragment, corrected_fragment]).issubset({'的', '地', '得'}))
            is_blacklisted = (original_fragment.lower() == corrected_fragment.lower() or re.match(r'^[A-Z]+$', original_fragment) or original_fragment in ['AI', 'Python'])
            if is_whitelisted and not is_blacklisted:
                final_parts.append(corrected_fragment)
            else:
                final_parts.append(original_fragment)
                log_stage = stage_map.get(stage, stage)
                log_tag = tag_map.get(tag, tag)
                logs_list.append(
                    f"【审查阶段】: {log_stage}\n"
                    f"【操作类型】: {log_tag}\n"
                    f"【原始片段】: '{original_fragment}'\n"
                    f"【模型建议】: '{corrected_fragment}' (已拒绝)"
                )
        return "".join(final_parts)

    def _get_highlight_details(self, original_sent, corrected_sent, offset, type_label):
        details = []
        matcher = difflib.SequenceMatcher(None, original_sent, corrected_sent)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal': continue
            original_fragment, corrected_fragment = original_sent[i1:i2], corrected_sent[j1:j2]
            label = ""
            if tag == 'delete': label = f"[{type_label}] 建议删除: '{original_fragment}'"
            elif tag == 'insert': label = f"[{type_label}] 建议添加: '{corrected_fragment}'"
            else: label = f"[{type_label}] {original_fragment} -> {corrected_fragment}"
            if label:
                details.append({"entity": label, "start": offset + i1, "end": offset + i2})
        return details


# ----------------- FastAPI 应用设置 -----------------
app = FastAPI(
    title="智能文本校对API",
    description="一个提供中文文本和文档（txt, docx, pdf）校对功能的API服务。",
    version="1.0.0",
)

# 实例化核心服务
correction_service = CorrectionService()

# API请求的数据模型
class TextRequest(BaseModel):
    text: str

# API端点
@app.post("/api/correct/text", summary="校对纯文本")
async def handle_text_correction(request: TextRequest) -> Dict[str, Any]:
    """
    接收一段纯文本，返回校对后的文本、高亮信息和审查日志。
    """
    if not request.text or not request.text.strip():
        raise HTTPException(status_code=400, detail="输入的 'text' 字段不能为空。")
    try:
        highlight_output, final_text, logs_output = correction_service.correct_text_pipeline(request.text)
        return {
            "corrected_text": final_text,
            "highlight_details": highlight_output,
            "logs": logs_output
        }
    except Exception as e:
        print(f"处理文本时发生错误: {e}")
        raise HTTPException(status_code=500, detail=f"服务器内部错误: {e}")

@app.post("/api/correct/document", summary="校对文档文件")
async def handle_document_correction(file: UploadFile = File(...)) -> Dict[str, Any]:
    """
    接收一个文档文件 (txt, docx, pdf)，返回校对后的纯文本、高亮信息、
    审查日志、OCR调试日志（仅PDF）以及Base64编码的校对后Word文档。
    """
    allowed_extensions = {".txt", ".docx", ".pdf"}
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型 '{file_ext}'。允许的类型为: {', '.join(allowed_extensions)}")

    temp_filepath = ""
    try:
        # 将上传文件保存到临时文件，以便于不同库的读取
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_filepath = temp_file.name

        final_text, highlight_result, doc_bytes, debug_ocr, logs_output = correction_service.correct_document_pipeline(temp_filepath, file.filename)

        # 将生成的docx文件内容进行Base64编码
        doc_base64 = base64.b64encode(doc_bytes).decode('utf-8')

        return {
            "original_filename": file.filename,
            "corrected_text": final_text,
            "highlight_details": highlight_result,
            "logs": logs_output,
            "ocr_debug_log": debug_ocr,
            "corrected_docx_base64": doc_base64
        }
    except Exception as e:
        print(f"处理文档时发生错误: {e}")
        raise HTTPException(status_code=500, detail=f"服务器内部错误: {e}")
    finally:
        # 确保临时文件被删除
        if temp_filepath and os.path.exists(temp_filepath):
            os.remove(temp_filepath)

@app.get("/", summary="API健康检查")
def read_root():
    return {"status": "ok", "message": "智能文本校对API服务已启动"}

# ----------------- 程序主入口 -----------------
if __name__ == '__main__':
    import uvicorn
    print("--- 启动API服务 ---")
    print("--- 模型已预加载，服务准备就绪 ---")
    print("--- 访问 http://127.0.0.1:8000/docs 查看可交互的API文档 ---")
    # 建议在生产环境中使用Gunicorn等多进程管理器来运行
    uvicorn.run(app, host="0.0.0.0", port=8000)